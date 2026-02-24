import subprocess
from docx import Document

MODEL = "mistral"

def gerarQuestoes(disciplina, assunto, qtdQuestoes, nivelQuestoes):
    prompt = f"""
    Você é um gerador profissional de questões educacionais.

Regras:
- Gere EXATAMENTE {qtdQuestoes} questão(ões).
- NÃO gere mais.
- NÃO gere menos.
- Responda SOMENTE com as questões.
- NÃO escreva introdução.
- NÃO escreva explicações.
- NÃO escreva respostas.

Formato:
1. Pergunta


Disciplina: {disciplina}
Assunto: {assunto}
Nível: {nivelQuestoes}
Idioma: Português
"""
    comando = ["ollama", "run", MODEL, prompt]
    questoes = subprocess.run(
        comando,
        capture_output=True,
        text=True,
        encoding="utf-8"
    )
    if questoes.returncode != 0:
        print("Erro:", questoes.stderr)
        return "Erro ao gerar questões."

    return questoes.stdout.strip()
def salvar_word(nome_arquivo, conteudo):
    doc = Document()
    doc.add_heading("Lista de Exercícios", level=1) 
    for linha in conteudo.splitlines():
        doc.add_paragraph(linha) 

    doc.save(nome_arquivo)


print("="*45)
print("📘 Gerador de Questões com IA")
print("Digite 'sair' para encerrar")
print("="*45)

while True:
    disciplina = input("Digite a disciplina: ")
    if disciplina.lower() == "sair":
        break
    assunto = input("Digite o assunto:")
    while True:
        try:
            qtdQuestoes = int(input("Quantidade: "))
            break
        except:
            print("Digite apenas um número.")

    nivelQuestoes = input("Digite o nível das questões: ")
    print("Gerando as questões...")

    resposta = gerarQuestoes(disciplina, assunto, qtdQuestoes, nivelQuestoes)

    salvar_word("exercicios.docx", resposta)
    print("📄 Arquivo Word criado: exercicios.docx")
    print(resposta)